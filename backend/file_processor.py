import pdfplumber
from docx import Document
import os
from typing import Dict, Any
import asyncio
from concurrent.futures import ThreadPoolExecutor

class FileProcessor:
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=4)
    
    async def extract_content(self, file_path: str, content_type: str) -> Dict[str, Any]:
        """
        异步提取文件内容
        """
        if content_type == "application/pdf":
            return await self._extract_pdf_content(file_path)
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            return await self._extract_word_content(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {content_type}")
    
    async def _extract_pdf_content(self, file_path: str) -> Dict[str, Any]:
        """
        提取PDF文件内容
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, self._extract_pdf_sync, file_path)
    
    def _extract_pdf_sync(self, file_path: str) -> Dict[str, Any]:
        """
        同步提取PDF内容
        """
        content = ""
        metadata = {
            "pages": 0,
            "extraction_method": "pdfplumber"
        }
        
        try:
            # 使用pdfplumber提取PDF内容
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        content += f"\n--- 第 {page_num} 页 ---\n"
                        content += page_text + "\n"
                
                # 获取PDF元数据
                if pdf.metadata:
                    metadata.update({
                        "title": pdf.metadata.get("Title", ""),
                        "author": pdf.metadata.get("Author", ""),
                        "creator": pdf.metadata.get("Creator", ""),
                        "creation_date": str(pdf.metadata.get("CreationDate", "")),
                    })
        
        except Exception as e:
            raise Exception(f"PDF提取失败: {str(e)}")
        
        if not content.strip():
            raise Exception("PDF文件中未找到可提取的文本内容")
        
        return {
            "content": content.strip(),
            "metadata": metadata
        }
    
    async def _extract_word_content(self, file_path: str) -> Dict[str, Any]:
        """
        提取Word文档内容
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, self._extract_word_sync, file_path)
    
    def _extract_word_sync(self, file_path: str) -> Dict[str, Any]:
        """
        同步提取Word内容
        """
        try:
            doc = Document(file_path)
            
            content = ""
            paragraph_count = 0
            
            # 提取段落内容
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n\n"
                    paragraph_count += 1
            
            # 提取表格内容
            table_count = len(doc.tables)
            if table_count > 0:
                content += "\n--- 表格内容 ---\n"
                for table_num, table in enumerate(doc.tables, 1):
                    content += f"\n表格 {table_num}:\n"
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        content += " | ".join(row_data) + "\n"
                    content += "\n"
            
            metadata = {
                "paragraphs": paragraph_count,
                "tables": table_count,
                "extraction_method": "python-docx"
            }
            
            # 尝试获取文档属性
            try:
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "creator": core_props.author or "",
                    "creation_date": str(core_props.created) if core_props.created else "",
                    "modified_date": str(core_props.modified) if core_props.modified else "",
                })
            except:
                pass
            
            if not content.strip():
                raise Exception("Word文档中未找到可提取的内容")
            
            return {
                "content": content.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            raise Exception(f"Word文档提取失败: {str(e)}")
